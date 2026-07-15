"""
Generate SVCC_vs_Actuated_Table.docx with significance levels (* p<0.05, ** p<0.01).
Paired t-test across common seeds (SVCC vs Actuated) for each cell.
"""

import os, re
import numpy as np
from scipy import stats
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE     = os.path.dirname(os.path.abspath(__file__))
RES_BASE = os.path.join(BASE, "Results")

# ── data directories & file patterns ──────────────────────────────────────────
INTERSECTION_TYPES = [
    {
        "label":       "Concurrent\nPermitted Right",
        "dir":         "Concurrent-Permitted Right",
        "act_pat":     r"actuated_penetration\(1\)_EVratio\(0\.5\)_Concurrent_permitted_right_Asymmetric_{ped}_seed\((\d+)\)\.txt$",
        "svcc_pat":    r"multi_scale_penetration\(1\)_EVratio\(0\.5\)_Concurrent_permitted_right_Asymmetric_\(1, 0, 0(?:\.0)?\)_{ped}_seed\((\d+)\)\.txt$",
    },
    {
        "label":       "Concurrent\nProtected Right",
        "dir":         "Concurrent-Protected Right",
        "act_pat":     r"actuated_penetration\(1\)_EVratio\(0\.5\)_Concurrent_protected_right_Asymmetric_{ped}_seed\((\d+)\)\.txt$",
        "svcc_pat":    r"multi_scale_penetration\(1\)_EVratio\(0\.5\)_Concurrent_protected_right_Asymmetric_\(1, 0, 0(?:\.0)?\)_{ped}_seed\((\d+)\)\.txt$",
    },
    {
        "label":       "Concurrent\nLPI-LBI",
        "dir":         "Concurrent-LPI",
        "act_pat":     r"actuated_penetration\(1\)_EVratio\(0\.5\)_Concurrent_LPI_Asymmetric_{ped}_seed\((\d+)\)\.txt$",
        "svcc_pat":    r"multi_scale_penetration\(1\)_EVratio\(0\.5\)_Concurrent_LPI_Asymmetric_\(1, 0, 0(?:\.0)?\)_{ped}_seed\((\d+)\)\.txt$",
    },
    {
        "label":       "Concurrent\nDelayed Turn",
        "dir":         "Concurrent-delayed turn",
        "act_pat":     r"actuated_penetration\(1\)_EVratio\(0\.5\)_Concurrent_delayed_turn_Asymmetric_{ped}_seed\((\d+)\)\.txt$",
        "svcc_pat":    r"multi_scale_penetration\(1\)_EVratio\(0\.5\)_Concurrent_delayed_turn_Asymmetric_\(1, 0, 0(?:\.0)?\)_{ped}_seed\((\d+)\)\.txt$",
    },
    {
        "label":       "Exclusive",
        "dir":         None,
        "act_pat":     None,
        "svcc_pat":    None,
    },
]

PED_LEVELS = ["LowPed", "MedPed", "HighPed"]
PED_LABEL  = {"LowPed": "Low Demand", "MedPed": "Medium Demand", "HighPed": "High Demand"}

METRICS = [
    # (category, metric_label, parse_key)
    ("Sustainability", "Fuel Consumption (mg/veh)",        "fuel"),
    ("Sustainability", "Power Consumption (KW.s/veh)",     "power"),
    ("Mobility",       "Vehicle Waiting Time (s/veh)",     "waiting_time"),
    ("Mobility",       "Vehicle Time Loss (s/veh)",        "time_loss"),
    ("Mobility",       "Vehicle Queue Length (m)",         "queue_length"),
    ("Mobility",       "Bike Waiting Time (s/bike)",       "bike_wait"),
    ("Mobility",       "Bike Time Loss (s/bike)",          "bike_timeloss"),
    ("Mobility",       "Bike Queue Length (m)",            "bike_queue"),
    ("Mobility",       "Pedestrian Time Loss (s/ped)",     "ped_timeloss"),
    ("Safety",         "# Veh-Ped Conflicts",              "veh_ped_conf"),
    ("Safety",         "# Veh-Bike Conflicts",             "veh_bike_conf"),
    ("Safety",         "# Veh-Veh Conflicts",              "veh_veh_conf"),
]

PARSE_PATTERNS = {
    "fuel":          r"average fuel consumption \(external model\) for CAVs (?:multi_scale|actuated) \(in mg\): ([\d.]+)",
    "power":         r"average power consumption \(external model\) for CAVs (?:multi_scale|actuated) \(in Kw\.s\): ([\d.]+)",
    "waiting_time":  r"average waiting time for (?:multi_scale|actuated) scenario \(in s\): ([\d.]+)",
    "time_loss":     r"average time loss for (?:multi_scale|actuated) scenario \(in s\): ([\d.]+)",
    "queue_length":  r"average queue length for (?:multi_scale|actuated) scenario \(in m\): ([\d.]+)",
    "bike_wait":     r"average bike \(sep\+int\) waiting time for (?:multi_scale|actuated) scenario \(in s\): ([\d.]+)",
    "bike_timeloss": r"average bike \(sep\+int\) time loss for (?:multi_scale|actuated) scenario \(in s\): ([\d.]+)",
    "bike_queue":    r"average bike \(sep\) queue length for (?:multi_scale|actuated) scenario \(in m\): ([\d.]+)",
    "ped_timeloss":  r"average pedestrian time loss for (?:multi_scale|actuated) scenario \(in s\): ([\d.]+)",
    "veh_ped_conf":  r"number of right-turn conflicts between vehicles and pedestrians: ([\d.]+)",
    "veh_bike_conf": r"number of conflicts between vehicles and cyclists: ([\d.]+)",
    "veh_veh_conf":  r"number of conflicts between vehicles and vehicles: ([\d.]+)",
}

# ── parsing ────────────────────────────────────────────────────────────────────

def parse_file(path):
    with open(path) as f:
        text = f.read()
    return {k: float(m.group(1)) if (m := re.search(p, text)) else None
            for k, p in PARSE_PATTERNS.items()}


def load_per_seed(directory, pattern_template, ped):
    pat = re.compile(pattern_template.replace("{ped}", ped))
    result = {}
    for f in os.listdir(directory):
        m = pat.match(f)
        if m:
            result[int(m.group(1))] = parse_file(os.path.join(directory, f))
    return result


def mean_over_common_seeds(seed_dict_a, seed_dict_b, metric_key):
    """Mean of seed_dict_a and seed_dict_b restricted to common seeds for metric_key."""
    common = set(seed_dict_a) & set(seed_dict_b)
    a_vals, b_vals = [], []
    for s in sorted(common):
        av = seed_dict_a[s].get(metric_key) if seed_dict_a.get(s) else None
        bv = seed_dict_b[s].get(metric_key) if seed_dict_b.get(s) else None
        if av is not None and bv is not None:
            a_vals.append(av)
            b_vals.append(bv)
    a_mean = float(np.mean(a_vals)) if a_vals else None
    b_mean = float(np.mean(b_vals)) if b_vals else None
    return a_mean, b_mean


def paired_ttest(act_seed_dict, svcc_seed_dict, metric_key):
    """Paired t-test SVCC vs Actuated across common seeds. Returns '' / '*' / '**'."""
    common = set(act_seed_dict) & set(svcc_seed_dict)
    a, b = [], []
    for s in sorted(common):
        av = act_seed_dict[s].get(metric_key) if act_seed_dict[s] else None
        bv = svcc_seed_dict[s].get(metric_key) if svcc_seed_dict[s] else None
        if av is not None and bv is not None:
            a.append(av)
            b.append(bv)
    if len(a) < 2:
        return ""
    _, p = stats.ttest_rel(b, a)
    if p < 0.01:
        return "**"
    if p < 0.05:
        return "*"
    return ""

# ── build cell data ────────────────────────────────────────────────────────────

def build_data():
    """
    Returns data[ped][itype_idx] = {
        'act_mean': {metric: val},
        'svcc_mean': {metric: val},
        'sig': {metric: '' | '*' | '**'},
    }
    """
    data = {}
    for ped in PED_LEVELS:
        data[ped] = []
        for itype in INTERSECTION_TYPES:
            if itype["dir"] is None:
                data[ped].append(None)
                continue
            directory = os.path.join(RES_BASE, itype["dir"])
            act_seed  = load_per_seed(directory, itype["act_pat"],  ped)
            svcc_seed = load_per_seed(directory, itype["svcc_pat"], ped)
            act_mean  = {}
            svcc_mean = {}
            sig = {}
            for _, _, mk in METRICS:
                a_m, b_m = mean_over_common_seeds(act_seed, svcc_seed, mk)
                act_mean[mk]  = a_m
                svcc_mean[mk] = b_m
                sig[mk] = paired_ttest(act_seed, svcc_seed, mk)
            data[ped].append({
                "act_mean":  act_mean,
                "svcc_mean": svcc_mean,
                "sig":       sig,
                "n_act":     len(act_seed),
                "n_svcc":    len(svcc_seed),
            })
    return data

# ── Word document generation ───────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val is not None:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val", "single"))
            el.set(qn("w:sz"),    val.get("sz",  "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "000000"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_run(para, text, bold=False, color=None, size=8):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run


def cell_para(cell, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    para = cell.paragraphs[0]
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return para


def generate_docx(data):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(1.2)
        section.right_margin  = Cm(1.2)

    # ── table layout ──────────────────────────────────────────────────────────
    # Columns:
    #   0: Demand level (merged vertically per demand)
    #   1: Category (merged vertically per category per demand)
    #   2: Metric label
    #   3,4: Concurrent Permitted Right (Actuated, SVCC)
    #   5,6: Concurrent Protected Right (Actuated, SVCC)
    #   7,8: Concurrent LPI-LBI (Actuated, SVCC)
    #   9,10: Concurrent Delayed Turn (Actuated, SVCC)
    #   11: Exclusive (SVCC only)

    n_cols = 12
    # Header rows: 3 rows for column headers
    # Data rows: 3 demand levels × 12 metrics = 36 rows
    header_rows = 3
    data_rows   = len(PED_LEVELS) * len(METRICS)
    total_rows  = header_rows + data_rows

    table = doc.add_table(rows=total_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Column widths (cm)
    col_widths = [1.8, 1.6, 3.8,  # demand, category, metric
                  2.0, 2.4,        # Perm Right Actuated, SVCC
                  2.0, 2.4,        # Prot Right Actuated, SVCC
                  2.0, 2.4,        # LPI Actuated, SVCC
                  2.0, 2.4,        # Delayed Actuated, SVCC
                  2.0]             # Exclusive SVCC
    for i, w in enumerate(col_widths):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    # ── HEADER ROW 0: Big intersection-type labels ─────────────────────────
    row0 = table.rows[0].cells

    # First 3 cols blank
    for c in range(3):
        set_cell_bg(row0[c], "2C3E50")
        row0[c].merge(table.rows[1].cells[c]).merge(table.rows[2].cells[c])

    para = cell_para(row0[0])
    add_run(para, "Metrics", bold=True, color="FFFFFF", size=9)

    # Concurrent headers (merged 2 cols each across rows 0-1)
    itype_labels = [
        ("Concurrent\nPermitted Right", 3),
        ("Concurrent\nProtected Right", 5),
        ("Concurrent\nLPI-LBI",         7),
        ("Concurrent\nDelayed Turn",     9),
        ("Exclusive",                   11),
    ]
    for label, col in itype_labels:
        cell = row0[col]
        if col < 11:
            # Merge two cols for Concurrent types (rows 0-1)
            cell = row0[col].merge(row0[col + 1])
        # Merge rows 0-1 for Exclusive (single col)
        if col == 11:
            cell = row0[11].merge(table.rows[1].cells[11])
        set_cell_bg(cell, "1A5276")
        para = cell_para(cell)
        add_run(para, label, bold=True, color="FFFFFF", size=9)

    # ── HEADER ROW 1: Sub-headers (Actuated / SVCC per type) ──────────────
    row1 = table.rows[1].cells
    for itype_idx, (_, col) in enumerate([(None, 3), (None, 5), (None, 7), (None, 9)]):
        for offset, sublabel in [(0, "Actuated\n(Baseline)"), (1, "SVCC")]:
            c = col + offset
            set_cell_bg(row1[c], "2980B9" if offset == 1 else "566573")
            para = cell_para(row1[c])
            add_run(para, sublabel, bold=True, color="FFFFFF", size=8)
        # row1[11] already merged up in row0 processing

    # ── HEADER ROW 2: Demand / Category / Metric labels ───────────────────
    row2 = table.rows[2].cells
    for c in range(3):
        set_cell_bg(row2[c], "2C3E50")  # will be covered by data rows

    # Add demand / category / metric column sub-headers in row 2
    for c in range(3):
        set_cell_bg(row2[c], "273746")
    sub_headers = ["Demand", "Category", "Metric"]
    for c, sh in enumerate(sub_headers):
        para = cell_para(row2[c])
        add_run(para, sh, bold=True, color="FFFFFF", size=8)

    # Col headers for data cols (row 2)
    for c in range(3, n_cols):
        set_cell_bg(row2[c], "273746")

    # ── DATA ROWS ──────────────────────────────────────────────────────────
    # Colors
    DEMAND_BG   = {"LowPed": "EBF5FB", "MedPed": "E8F8F5", "HighPed": "FEF9E7"}
    DEMAND_HDR  = {"LowPed": "2471A3", "MedPed": "1E8449", "HighPed": "B7950B"}
    CAT_BG      = {"Sustainability": "D6EAF8", "Mobility": "D5F5E3", "Safety": "FDEBD0"}
    ACT_BG      = "D5D8DC"
    SVCC_BG     = "FDFEFE"
    SVCC_SIG_BG = "EBF5FB"

    row_idx = header_rows
    for ped in PED_LEVELS:
        ped_start = row_idx
        categories_seen = {}
        cat_starts = {}

        for m_idx, (cat, metric_label, mk) in enumerate(METRICS):
            row = table.rows[row_idx].cells

            # Track category spans
            if cat not in categories_seen:
                categories_seen[cat] = row_idx
                cat_starts[cat] = row_idx

            # ── Demand cell (col 0) ──────────────────────────────────────
            set_cell_bg(row[0], DEMAND_BG[ped])
            if m_idx == 0:
                para = cell_para(row[0])
                add_run(para, PED_LABEL[ped], bold=True,
                        color=DEMAND_HDR[ped].lstrip("#") if not DEMAND_HDR[ped].startswith("#") else DEMAND_HDR[ped][1:],
                        size=8)

            # ── Category cell (col 1) ────────────────────────────────────
            set_cell_bg(row[1], CAT_BG[cat])
            para = cell_para(row[1])
            add_run(para, cat, bold=True, size=8)

            # ── Metric label (col 2) ─────────────────────────────────────
            set_cell_bg(row[2], "FDFEFE")
            para = cell_para(row[2], WD_ALIGN_PARAGRAPH.LEFT)
            add_run(para, metric_label, size=8)

            # ── Data cells ───────────────────────────────────────────────
            for itype_idx, itype in enumerate(INTERSECTION_TYPES):
                idata = data[ped][itype_idx]

                if idata is None:
                    # Exclusive — N/A
                    cell = row[11]
                    set_cell_bg(cell, "EAECEE")
                    para = cell_para(cell)
                    add_run(para, "N/A", color="7F8C8D", size=8)
                    continue

                act_col  = 3 + itype_idx * 2
                svcc_col = 4 + itype_idx * 2

                act_val  = idata["act_mean"].get(mk)
                svcc_val = idata["svcc_mean"].get(mk)
                sig      = idata["sig"].get(mk, "")

                # Actuated cell
                act_cell = row[act_col]
                set_cell_bg(act_cell, ACT_BG)
                para = cell_para(act_cell)
                txt = f"{act_val:.2f}" if act_val is not None else "N/A"
                add_run(para, txt, size=8)

                # SVCC cell
                svcc_cell = row[svcc_col]
                if svcc_val is not None and act_val is not None:
                    pct  = (svcc_val - act_val) / act_val * 100
                    sign = "+" if pct >= 0 else ""
                    # Color: green tint if improvement (lower), red tint if worse
                    # For conflicts and most metrics: lower = better
                    if pct < 0:
                        bg = "D5F5E3"  # green tint = improvement
                    elif pct > 0:
                        bg = "FADBD8"  # red tint = regression
                    else:
                        bg = SVCC_BG
                    set_cell_bg(svcc_cell, bg)
                    para = cell_para(svcc_cell)
                    val_str = f"{svcc_val:.2f}"
                    pct_str = f"({sign}{pct:.1f}%)"
                    if sig:
                        add_run(para, val_str + sig, bold=True, size=8)
                        para.add_run("\n")
                        r2 = para.add_run(pct_str)
                        r2.font.size = Pt(7.5)
                        r2.bold = True
                    else:
                        add_run(para, val_str, size=8)
                        para.add_run("\n")
                        r2 = para.add_run(pct_str)
                        r2.font.size = Pt(7.5)
                else:
                    set_cell_bg(svcc_cell, SVCC_BG)
                    para = cell_para(svcc_cell)
                    txt = f"{svcc_val:.2f}" if svcc_val is not None else "N/A"
                    add_run(para, txt, size=8)

            row_idx += 1

        # Merge demand column cells (col 0) for this ped block
        if len(PED_LEVELS) > 1 and (row_idx - ped_start) > 1:
            merged = table.rows[ped_start].cells[0]
            for r in range(ped_start + 1, row_idx):
                merged = merged.merge(table.rows[r].cells[0])
            set_cell_bg(merged, DEMAND_BG[ped])
            merged.paragraphs[0].clear()
            para = cell_para(merged)
            add_run(para, PED_LABEL[ped], bold=True,
                    color=DEMAND_HDR[ped],
                    size=9)

        # Merge category column cells (col 1) for each category block
        cat_order = []
        for cat, _, _ in METRICS:
            if cat not in [c for c, _ in cat_order]:
                cat_order.append((cat, ped_start + METRICS.index(next(m for m in METRICS if m[0] == cat))))

        cat_groups = {}
        for m_idx, (cat, _, _) in enumerate(METRICS):
            cat_groups.setdefault(cat, []).append(ped_start + m_idx)

        for cat, row_indices in cat_groups.items():
            if len(row_indices) > 1:
                merged = table.rows[row_indices[0]].cells[1]
                for r in row_indices[1:]:
                    merged = merged.merge(table.rows[r].cells[1])
                set_cell_bg(merged, CAT_BG[cat])
                merged.paragraphs[0].clear()
                para = cell_para(merged)
                add_run(para, cat, bold=True, size=8)

    # ── add footnote ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run(
        "* p < 0.05  ** p < 0.01  (paired t-test, SVCC vs Actuated across simulation seeds).\n"
        "Green shading = improvement over Actuated baseline; Red shading = regression. "
        "Bold values indicate statistically significant differences."
    )
    run.font.size = Pt(8)
    run.italic    = True

    out_path = os.path.join(RES_BASE, "SVCC_vs_Actuated_Table.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


# ── main ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Loading and computing data...")
    data = build_data()

    # Print summary of significance findings
    for ped in PED_LEVELS:
        print(f"\n{PED_LABEL[ped]}:")
        for itype_idx, itype in enumerate(INTERSECTION_TYPES):
            idata = data[ped][itype_idx]
            if idata is None:
                continue
            sigs = {mk: idata["sig"][mk] for _, _, mk in METRICS if idata["sig"][mk]}
            print(f"  {itype['label'].replace(chr(10), ' ')}: "
                  f"n_act={idata['n_act']}, n_svcc={idata['n_svcc']}, "
                  f"significant metrics: {sigs}")

    print("\nGenerating Word document...")
    out = generate_docx(data)
    print("Done.")
